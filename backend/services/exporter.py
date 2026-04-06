# backend/services/exporter.py
"""
Export a written research paper to DOCX or PDF.
"""
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.lib import colors


SECTION_ORDER = ["abstract", "introduction", "related_work", "methodology", "results", "discussion", "conclusion", "references"]
SECTION_LABELS = {
    "abstract": "Abstract",
    "introduction": "1. Introduction",
    "related_work": "2. Related Work",
    "methodology": "3. Methodology",
    "results": "4. Results",
    "discussion": "5. Discussion",
    "conclusion": "6. Conclusion",
    "references": "References",
}


def export_docx(title: str, authors: str, sections: dict) -> bytes:
    """Generate a formatted DOCX file. Returns bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    # Authors
    if authors:
        auth_para = doc.add_paragraph()
        auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = auth_para.add_run(authors)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()  # spacer

    for sec_id in SECTION_ORDER:
        content = sections.get(sec_id, "")
        if not content:
            continue

        # Section heading
        heading = doc.add_heading(SECTION_LABELS[sec_id], level=1)
        heading.runs[0].font.size = Pt(13)
        heading.runs[0].font.bold = True

        # Body — strip markdown bold markers for clean docx
        clean = content.replace("**", "")
        for paragraph_text in clean.split("\n"):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                p.style.font.size = Pt(11)

        doc.add_paragraph()  # spacer between sections

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_pdf(title: str, authors: str, sections: dict) -> bytes:
    """Generate a formatted PDF. Returns bytes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=72,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontSize=18, spaceAfter=6)
    author_style = ParagraphStyle("Author", parent=styles["Normal"], fontSize=12, textColor=colors.grey, spaceAfter=20, alignment=1)
    heading_style = ParagraphStyle("Heading", parent=styles["Heading1"], fontSize=13, spaceBefore=16, spaceAfter=6)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=6)

    story = [
        Paragraph(title, title_style),
        Paragraph(authors or "", author_style),
        Spacer(1, 0.2 * inch),
    ]

    for sec_id in SECTION_ORDER:
        content = sections.get(sec_id, "")
        if not content:
            continue
        story.append(Paragraph(SECTION_LABELS[sec_id], heading_style))
        clean = content.replace("**", "").replace("*", "")
        for para_text in clean.split("\n"):
            if para_text.strip():
                story.append(Paragraph(para_text.strip(), body_style))
        story.append(Spacer(1, 0.1 * inch))

    doc.build(story)
    return buffer.getvalue()
